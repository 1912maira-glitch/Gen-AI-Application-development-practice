````python
import io
import json
import os
import re
from typing import Any, Dict

import streamlit as st
from google import genai
from google.genai import types
from pypdf import PdfReader
from docx import Document


MODEL_NAME = "gemini-3.5-flash"

MAX_POINTS = {
    "keyword_match": 35,
    "experience_relevance": 25,
    "ats_structure": 15,
    "role_alignment": 10,
    "education_certifications": 5,
    "impact_action_language": 10,
}

st.set_page_config(
    page_title="Resume ATS Analyzer",
    page_icon="📄",
    layout="wide",
)


def get_api_key() -> str:
    """Read Gemini API key from Streamlit Secrets or environment."""
    try:
        if "GEMINI_API_KEY" in st.secrets:
            return str(st.secrets["GEMINI_API_KEY"]).strip()
    except Exception:
        pass

    return os.getenv("GEMINI_API_KEY", "").strip()


def extract_text(uploaded_file) -> str:
    """Extract text from PDF, DOCX, or TXT."""
    suffix = uploaded_file.name.lower().rsplit(".", 1)[-1]
    data = uploaded_file.getvalue()

    if suffix == "pdf":
        reader = PdfReader(io.BytesIO(data))
        pages = []

        for page in reader.pages:
            pages.append(page.extract_text() or "")

        return "\n".join(pages)

    if suffix == "docx":
        doc = Document(io.BytesIO(data))
        parts = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]

                if any(cells):
                    parts.append(" | ".join(cells))

        return "\n".join(parts)

    if suffix == "txt":
        return data.decode("utf-8", errors="ignore")

    raise ValueError(
        "Unsupported file type. Please upload PDF, DOCX, or TXT."
    )


def clean_json(text: str) -> Dict[str, Any]:
    """Parse JSON even if Gemini returns a markdown code fence."""
    text = text.strip()

    if text.startswith("```"):
        text = re.sub(
            r"^```(?:json)?\s*",
            "",
            text,
            flags=re.IGNORECASE,
        )
        text = re.sub(r"\s*```$", "", text)

    return json.loads(text)


def clamp(value: Any, minimum: int, maximum: int) -> int:
    """Safely clamp a numeric value."""
    try:
        value = int(value)
    except (TypeError, ValueError):
        value = minimum

    return max(minimum, min(maximum, value))


def validate_result(result: Dict[str, Any]) -> Dict[str, Any]:
    """
    Validate Gemini's response.

    The final score is calculated in Python from the category scores.
    Gemini cannot override the final score.
    """

    breakdown = result.get("score_breakdown")

    if not isinstance(breakdown, dict):
        breakdown = {}

    for key, maximum in MAX_POINTS.items():
        breakdown[key] = clamp(
            breakdown.get(key, 0),
            0,
            maximum,
        )

    result["score_breakdown"] = breakdown

    # Calculate final score ourselves.
    result["ats_score"] = sum(breakdown.values())

    list_fields = [
        "matched_keywords",
        "missing_keywords",
        "strengths",
        "ats_format_risks",
        "rewrites",
        "improvements",
    ]

    for field in list_fields:
        if not isinstance(result.get(field), list):
            result[field] = []

    # Keep only valid improvement objects.
    cleaned_improvements = []

    for item in result["improvements"]:
        if not isinstance(item, dict):
            continue

        priority = item.get("priority", "Medium")

        if priority not in {"High", "Medium", "Low"}:
            priority = "Medium"

        cleaned_improvements.append(
            {
                "priority": priority,
                "issue": str(item.get("issue", "")),
                "recommendation": str(
                    item.get("recommendation", "")
                ),
                "example": str(item.get("example", "")),
            }
        )

    result["improvements"] = cleaned_improvements

    return result


def analyze_resume(
    resume_text: str,
    job_description: str,
) -> Dict[str, Any]:

    api_key = get_api_key()

    if not api_key:
        raise RuntimeError(
            "GEMINI_API_KEY is not configured. Add it to "
            "Streamlit Secrets or set it as an environment variable."
        )

    client = genai.Client(api_key=api_key)

    prompt = f"""
You are an evidence-first ATS resume evaluator and career coach.

Your task is to compare the RESUME against the TARGET JOB DESCRIPTION.

The ATS score is a HEURISTIC estimate. It is NOT an official score
from a real ATS vendor.

==================================================
CRITICAL: FACTUAL ACCURACY RULES
==================================================

The RESUME is the ONLY source of truth for the candidate's experience.

NEVER invent or fabricate:

- percentages
- numbers
- accuracy
- latency
- users
- revenue
- money saved
- time saved
- performance improvements
- employers
- job titles
- responsibilities
- technologies
- certifications
- projects
- achievements
- outcomes
- dates

If a number does not appear in the resume, DO NOT create a number.

For example, if the resume says:

"Built an AI automation workflow"

DO NOT rewrite it as:

"Built an AI automation workflow that reduced processing time by 40%."

The 40% is fabricated.

Instead, write something like:

"Built an AI automation workflow using n8n for automated inquiry
classification and lead follow-up."

You may recommend that the candidate add a real metric IF they
actually measured one, but never invent the metric.

Use phrases such as:

"Add a real metric if you measured one."

==================================================
TECHNOLOGY ACCURACY
==================================================

Do not claim that a candidate knows a technology simply because it
appears in the job description.

For example:

If the JD says:
"Experience with Pydantic"

and the resume does not mention Pydantic, it is missing.

Do NOT suggest:

"Add Pydantic to your resume"

unless the candidate actually used Pydantic.

Instead say:

"Learn or gain hands-on experience with Pydantic if relevant to your
target roles."

==================================================
SEMANTIC MATCHING
==================================================

A skill may be considered matched when the resume uses a clearly
equivalent phrase.

Example:

Resume:
"API integration"

JD:
"API integrations"

This can count as a match.

However, do not assume a more specific technology from a generic term.

Example:

Resume:
"API integration"

JD:
"REST APIs"

Do NOT automatically claim REST APIs are present.

==================================================
STRUCTURED OUTPUTS
==================================================

Only consider "structured outputs" matched when the resume provides
evidence such as:

- JSON output
- structured LLM responses
- schemas
- typed objects
- Pydantic
- function/tool output
- structured response formats

Do not treat "structured learning plans" as automatically equivalent
to "structured LLM outputs."

==================================================
ATS FORMAT ANALYSIS
==================================================

The resume may have extraction artifacts caused by PDF/DOCX parsing.

Examples:

| text |
blank lines
repeated spaces
table extraction artifacts

Do NOT call these actual formatting problems unless the extracted
content provides reasonable evidence.

Because you are receiving extracted text rather than the original
visual document, clearly acknowledge this limitation when necessary.

==================================================
SCORING
==================================================

Use exactly these maximum scores:

Keyword / skill match: 35
Relevant experience / projects: 25
ATS structure: 15
Role alignment: 10
Education / certifications: 5
Impact / action language: 10

TOTAL: 100

Each category MUST be between zero and its maximum.

The final score is the sum of the six category scores.

==================================================
MATCHED KEYWORDS
==================================================

List important job-description skills that are actually supported
by the resume.

Prefer the terminology used by the resume.

==================================================
MISSING KEYWORDS
==================================================

List important job-description skills that are not supported by the
resume.

Do not list every minor word.

Do not penalize the candidate for ordinary words.

==================================================
IMPROVEMENTS
==================================================

Prioritize useful improvements.

Each improvement should contain:

1. The issue
2. A truthful recommendation
3. An example only if the example uses facts already present
   in the resume

If a metric would improve a bullet but there is no metric in the
resume, say:

"Add a real metric if you measured one."

Never create a sample percentage or number inside the resume rewrite.

==================================================
REWRITES
==================================================

Provide up to five bullet rewrites.

Every improved bullet MUST:

- preserve the original factual meaning
- use only facts from the resume
- improve clarity
- use strong action verbs where appropriate
- improve ATS keyword alignment when supported
- NOT introduce new technologies
- NOT introduce new responsibilities
- NOT introduce new achievements
- NOT introduce invented numbers

==================================================
OUTPUT
==================================================

Return ONLY valid JSON.

Use exactly this structure:

{{
  "ats_score": 0,

  "score_breakdown": {{
    "keyword_match": 0,
    "experience_relevance": 0,
    "ats_structure": 0,
    "role_alignment": 0,
    "education_certifications": 0,
    "impact_action_language": 0
  }},

  "summary": "string",

  "matched_keywords": [
    "string"
  ],

  "missing_keywords": [
    "string"
  ],

  "strengths": [
    "string"
  ],

  "improvements": [
    {{
      "priority": "High|Medium|Low",
      "issue": "string",
      "recommendation": "string",
      "example": "string"
    }}
  ],

  "ats_format_risks": [
    "string"
  ],

  "rewrites": [
    {{
      "original": "string",
      "improved": "string"
    }}
  ]
}}

==================================================
RESUME
==================================================

---BEGIN RESUME---

{resume_text[:50000]}

---END RESUME---

==================================================
TARGET JOB DESCRIPTION
==================================================

---BEGIN JOB DESCRIPTION---

{job_description[:20000]}

---END JOB DESCRIPTION---
"""

    response = client.models.generate_content(
        model=MODEL_NAME,
        contents=prompt,
        config=types.GenerateContentConfig(
            response_mime_type="application/json",
        ),
    )

    if not response.text:
        raise RuntimeError(
            "Gemini returned an empty response."
        )

    result = clean_json(response.text)

    return validate_result(result)


def display_score(score: int) -> None:
    st.metric(
        "Estimated ATS Score",
        f"{score}/100",
    )

    st.progress(score / 100)


def main() -> None:

    st.title("📄 Resume ATS Analyzer")

    st.caption(
        "Compare your resume with a target job description using "
        "Gemini-powered, evidence-based ATS analysis."
    )

    with st.sidebar:

        st.header("How it works")

        st.write(
            "1. Upload your resume.\n"
            "2. Paste the target job description.\n"
            "3. Gemini compares the resume and JD.\n"
            "4. Review matched skills and genuine gaps.\n"
            "5. Improve your resume using evidence-based suggestions."
        )

        st.info(
            "AI safety: the analyzer is instructed not to invent "
            "metrics, skills, experience, or achievements."
        )

        st.caption(
            "Privacy: extracted resume text is sent to the Gemini API "
            "for analysis."
        )

    uploaded_file = st.file_uploader(
        "Upload resume",
        type=["pdf", "docx", "txt"],
        help="Supported formats: PDF, DOCX, TXT.",
        max_upload_size=10,
    )

    job_description = st.text_area(
        "Target job description",
        height=260,
        placeholder=(
            "Paste the complete target job description here..."
        ),
    )

    if uploaded_file:

        try:

            resume_text = extract_text(uploaded_file)

            if not resume_text.strip():

                st.error(
                    "No text could be extracted. If this is a scanned "
                    "PDF, use an OCR-enabled PDF or upload DOCX/TXT."
                )

                return

            st.success(
                f"Extracted approximately "
                f"{len(resume_text):,} characters."
            )

            with st.expander(
                "Preview extracted resume text"
            ):
                st.text(resume_text[:5000])

        except Exception as exc:

            st.error(
                f"Could not read the resume: {exc}"
            )

            return

    else:
        resume_text = ""

    analyze_clicked = st.button(
        "Analyze Resume",
        type="primary",
        use_container_width=True,
        disabled=(
            not uploaded_file
            or not job_description.strip()
        ),
    )

    if analyze_clicked:

        with st.spinner(
            "Analyzing your resume with Gemini..."
        ):

            try:

                result = analyze_resume(
                    resume_text,
                    job_description.strip(),
                )

                st.session_state["analysis"] = result

            except Exception as exc:

                st.error(
                    f"Analysis failed: {exc}"
                )

    result = st.session_state.get("analysis")

    if not result:
        return

    st.divider()

    display_score(
        result["ats_score"]
    )

    st.caption(
        "This is an AI-generated heuristic estimate, not an official "
        "ATS score."
    )

    st.subheader("Summary")

    st.write(
        result.get(
            "summary",
            "No summary returned.",
        )
    )

    st.subheader("Score breakdown")

    breakdown = result.get(
        "score_breakdown",
        {},
    )

    cols = st.columns(3)

    labels = [
        (
            "Keyword match",
            "keyword_match",
        ),
        (
            "Experience relevance",
            "experience_relevance",
        ),
        (
            "ATS structure",
            "ats_structure",
        ),
        (
            "Role alignment",
            "role_alignment",
        ),
        (
            "Education/certifications",
            "education_certifications",
        ),
        (
            "Impact/action language",
            "impact_action_language",
        ),
    ]

    for index, (label, key) in enumerate(labels):

        with cols[index % 3]:

            st.metric(
                label,
                f"{breakdown.get(key, 0)}/"
                f"{MAX_POINTS[key]}",
            )

    left, right = st.columns(2)

    with left:

        st.subheader(
            "✅ Matched keywords"
        )

        matched = result.get(
            "matched_keywords",
            [],
        )

        if matched:
            st.write(
                ", ".join(matched)
            )
        else:
            st.write(
                "None identified."
            )

    with right:

        st.subheader(
            "⚠️ Missing / unconfirmed keywords"
        )

        missing = result.get(
            "missing_keywords",
            [],
        )

        if missing:
            st.write(
                ", ".join(missing)
            )
        else:
            st.write(
                "None identified."
            )

    st.subheader("Strengths")

    strengths = result.get(
        "strengths",
        [],
    )

    if strengths:

        for item in strengths:
            st.markdown(
                f"- {item}"
            )

    else:

        st.write(
            "No strengths returned."
        )

    st.subheader(
        "Priority improvements"
    )

    improvements = result.get(
        "improvements",
        [],
    )

    if improvements:

        for item in improvements:

            priority = item.get(
                "priority",
                "Medium",
            )

            st.markdown(
                f"**{priority}: "
                f"{item.get('issue', 'Improvement')}**"
            )

            recommendation = item.get(
                "recommendation",
                "",
            )

            if recommendation:
                st.write(
                    recommendation
                )

            example = item.get(
                "example",
                "",
            )

            if example:
                st.caption(
                    f"Example: {example}"
                )

    else:

        st.write(
            "No improvements returned."
        )

    st.subheader(
        "ATS format risks"
    )

    risks = result.get(
        "ats_format_risks",
        [],
    )

    if risks:

        for risk in risks:
            st.markdown(
                f"- {risk}"
            )

    else:

        st.write(
            "No major ATS format risks were identified from "
            "the extracted text."
        )

    rewrites = result.get(
        "rewrites",
        [],
    )

    if rewrites:

        st.subheader(
            "Suggested bullet rewrites"
        )

        st.warning(
            "Verify every rewrite against your actual experience. "
            "Do not add claims, technologies, or metrics that you "
            "cannot substantiate."
        )

        for item in rewrites:

            st.markdown(
                f"**Original:** "
                f"{item.get('original', '')}"
            )

            st.markdown(
                f"**Improved:** "
                f"{item.get('improved', '')}"
            )

            st.divider()

    st.download_button(
        "Download analysis as JSON",
        data=json.dumps(
            result,
            indent=2,
            ensure_ascii=False,
        ),
        file_name="resume_ats_analysis.json",
        mime="application/json",
        use_container_width=True,
    )


if __name__ == "__main__":
    main()
````
